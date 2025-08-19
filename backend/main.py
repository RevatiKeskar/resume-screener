from fastapi import FastAPI, UploadFile, File
import PyPDF2
import io
import docx
import re
from uuid import uuid4
import chromadb
from sentence_transformers import SentenceTransformer

app = FastAPI()
# load embedding model
model = SentenceTransformer("all-MiniLm-L6-v2")

# create a chromadb client  + collection
chroma_client = chromadb.PersistentClient(path = "resume_chroma_db")
collection = chroma_client.get_or_create_collection(name = "resumes")


@app.post("/upload_resume")
async def upload_resume(file: UploadFile = File(...)):
    # read content of the file
    file_content = await file.read()

    

    #extract text from all pages
    text = ""

    # Resume is pdf type case : 
    if file.filename.endswith(".pdf"):
        #convert bytes into pdf reader
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"


    # Resume is docx type case : 
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_content))
        for para in doc.paragraphs:
            text += para.text + "\n"

    # unsupported type
    else:
        return{"error" : "Unsupported file format."}
    
    # Finding out emails and phone no.s with regular expressions
    EMAIL_REGEX = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
    PHONE_REGEX = r"(?:\+?91[\s-]?)?[6-9]\d{9}"
    URL_REGEX   = r"(https?://\S+|www\.\S+)"

    emails = re.findall(EMAIL_REGEX, text)
    phones = re.findall(PHONE_REGEX, text)

    # deduping while keeping order
    emails = list(dict.fromkeys([e.strip() for e in emails]))
    phones = list(dict.fromkeys([p.strip() for p in phones]))

    def clean_text_for_llm():
        text_cleaned = re.sub(EMAIL_REGEX, " ", text)
        text_cleaned = re.sub(PHONE_REGEX, " ", text_cleaned)
        text_cleaned = re.sub(URL_REGEX, " ", text_cleaned)
        return text_cleaned
    text_cleaned = clean_text_for_llm()

    def store_resume_in_chroma(doc_text:str, metadata: dict):
        emb = model.encode([doc_text]).tolist()
        doc_id = f"resume_{uuid4()}"
        collection.add(
            documents = [doc_text],
            embeddings = emb,
            metadatas = [metadata],
            ids = [doc_id]
        )
        return doc_id
    
    emails_meta = emails[0] if emails else ""
    phones_meta = phones[0] if phones else ""
    doc_id = store_resume_in_chroma(text_cleaned, 
                                    {"filename":file.filename, "emails" : emails_meta, "phones" : phones_meta})
    
    return{
        "filename" : file.filename,
        "emails" : emails,
        "phones" : phones,
        "cleaned_text" : text_cleaned[:1000],
        "doc_id" : doc_id
    }

@app.post("/upload_job_description")
async def upload_job_description(
        file : UploadFile = File(...),
        top_k: int = 3
):
    content = await file.read()

    text = ""
    if file.filename.endswith(".pdf"):
        # Parse PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        for page in pdf_reader.pages:
            text += page.extract_text() or ""

    elif file.filename.endswith(".docx"):
        # Parse DOCX
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            text += para.text + "\n"

    else:
        # Fallback: assume plain text
        text = content.decode("utf-8", errors="ignore")

    # clean jd for embedding
    EMAIL_REGEX = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
    PHONE_REGEX = r"(?:\+?91[\s-]?)?[6-9]\d{9}"
    URL_REGEX   = r"(https?://\S+|www\.\S+)"

    def clean_text_for_llm(txt: str):
        txt = re.sub(EMAIL_REGEX, " ", txt)
        txt = re.sub(PHONE_REGEX, " ", txt)
        txt = re.sub(URL_REGEX, " ", txt)
        return txt
    
    cleaned_jd = clean_text_for_llm(text)

    # embed jd
    jd_embedding = model.encode([cleaned_jd]).tolist()

    # query top matching resumes
    results = collection.query(
        query_embeddings = jd_embedding,
        n_results = top_k
    )

    # format results
    matches = []
    for i in range(len(results["ids"][0])):
        matches.append({
            "doc_id" : results["ids"][0][i],
            "score" : results["distances"][0][i], # similary score
            "metadata" : results["metadatas"][0][i],
            "snippet" : results["documents"][0][i][:400] # preview
        })

    return {
        "job_description" : cleaned_jd[:500],
        "top_matches" : matches
    }



    